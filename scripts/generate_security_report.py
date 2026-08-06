from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
from datetime import date

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "entregables"
OUT.mkdir(exist_ok=True)
DOCX = OUT / "Reporte_Seguridad_Funnel_Comercial_CR3_2026-07-31.docx"

NAVY = "12305A"
BLUE = "2563EB"
LIGHT_BLUE = "EAF2FF"
PALE = "F5F8FC"
GREEN = "15803D"
GREEN_BG = "DCFCE7"
AMBER = "B45309"
AMBER_BG = "FEF3C7"
RED = "B91C1C"
RED_BG = "FEE2E2"
GRAY = "64748B"
WHITE = "FFFFFF"
BLACK = "0F172A"


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=BLACK, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    margins(cell)


def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def keep_row(row):
    trPr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)


def add_table(doc, headers, rows, widths=None, font=8.0):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, True, WHITE, 8.0)
        shade(hdr.cells[i], NAVY)
        if widths: hdr.cells[i].width = Inches(widths[i])
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        keep_row(table.rows[-1])
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, False, BLACK, font)
            shade(cells[i], WHITE if ridx % 2 == 0 else PALE)
            if widths: cells[i].width = Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title, body, level="info"):
    colors = {"high": (RED, RED_BG), "medium": (AMBER, AMBER_BG), "ok": (GREEN, GREEN_BG), "info": (BLUE, LIGHT_BLUE)}
    accent, bg = colors[level]
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(0.08)
    t.columns[1].width = Inches(6.35)
    shade(t.cell(0, 0), accent)
    shade(t.cell(0, 1), bg)
    margins(t.cell(0, 1), 120, 160, 120, 160)
    p = t.cell(0, 1).paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(accent)
    p2 = t.cell(0, 1).add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.08
    for r in p2.runs: r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix); r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.top_margin = Inches(0.72); sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(0.78); sec.right_margin = Inches(0.78)

styles = doc.styles
styles["Normal"].font.name = "Calibri"; styles["Normal"].font.size = Pt(10)
styles["Normal"].font.color.rgb = RGBColor.from_string(BLACK)
for name, size, color in (("Title", 28, NAVY), ("Heading 1", 17, NAVY), ("Heading 2", 13, BLUE), ("Heading 3", 10.5, NAVY)):
    st = styles[name]; st.font.name = "Calibri"; st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(10); st.paragraph_format.space_after = Pt(5)

# Header/footer
header = sec.header
hp = header.paragraphs[0]
hp.text = "GRUPO CR3  |  FUNNEL DE INTELIGENCIA COMERCIAL"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs: r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(GRAY)
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("CONFIDENCIAL - USO INTERNO   |   Evaluacion estatica al 31/julio/2026   |   ")
run.font.size = Pt(8); run.font.color.rgb = RGBColor.from_string(GRAY)
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
fp._p.append(fld)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(70)
r = p.add_run("GRUPO CR3")
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor.from_string(BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
r = p.add_run("Reporte integral de seguridad")
r.bold = True; r.font.size = Pt(30); r.font.color.rgb = RGBColor.from_string(NAVY)
p = doc.add_paragraph()
r = p.add_run("Plataforma Funnel de Inteligencia Comercial")
r.font.size = Pt(17); r.font.color.rgb = RGBColor.from_string(BLUE)
doc.add_paragraph("Evaluacion de arquitectura, codigo, controles, riesgos residuales y plan de tratamiento")
doc.add_paragraph()
add_callout(doc, "Clasificacion", "CONFIDENCIAL - USO INTERNO. Este documento contiene detalles de arquitectura y seguridad. No debe compartirse fuera de personal autorizado.", "high")
meta = add_table(doc, ["Dato", "Valor"], [
    ["Fecha de evaluacion", "31 de julio de 2026"],
    ["Version evaluada", "Repositorio vercel-supabase (estado local al corte)"],
    ["Ambiente objetivo", "Vercel + Supabase; futura migracion a servidor interno"],
    ["Tipo de revision", "Revision estatica de codigo y configuracion; no es pentest"],
    ["Responsable del documento", "Revision tecnica asistida por Codex"],
], [1.75, 4.65], 9)
doc.add_paragraph()
p = doc.add_paragraph("DICTAMEN EJECUTIVO")
p.style = styles["Heading 2"]
para(doc, "La plataforma cuenta con una base de seguridad superior a un prototipo comun: autenticacion Supabase, MFA para perfiles privilegiados, CAPTCHA, autorizacion de registros en servidor, cabeceras defensivas, auditoria y automatizacion de dependencias. Sin embargo, no debe considerarse plenamente endurecida hasta cerrar los riesgos altos detallados en este reporte.")
doc.add_page_break()

heading(doc, "1. Resumen ejecutivo", 1)
add_callout(doc, "Postura actual: riesgo residual ALTO", "La clasificacion no significa que la plataforma este comprometida. Significa que existen rutas de abuso relevantes y dependencias con vulnerabilidades altas que deben corregirse antes de declarar un nivel de seguridad empresarial maduro.", "high")
add_table(doc, ["Dimension", "Evaluacion", "Lectura ejecutiva"], [
    ["Identidad y acceso", "Parcialmente robusto", "MFA y CAPTCHA presentes; falta endurecer el alta automatica de perfiles y comprobar politicas operativas."],
    ["Autorizacion de datos", "Mixto", "Los registros principales se filtran en servidor; otros documentos compartidos aun pueden leerse completos."],
    ["Aplicacion y API", "Parcial", "Hay validaciones, limites y auditoria; falta una lista cerrada de claves y permisos por campo."],
    ["Base de datos", "Robusto con dependencia central", "RLS y privilegios directos revocados; la API usa service role y concentra el control."],
    ["Dependencias", "No aceptable al corte", "pnpm audit encontro 9 vulnerabilidades altas y 6 moderadas."],
    ["Continuidad", "Insuficiente", "No existe evidencia de una restauracion probada ni RTO/RPO aprobados."],
    ["Monitoreo", "Parcial", "Existe bitacora y webhook; falta demostrar entrega, retencion y respuesta operativa."],
], [1.35, 1.35, 3.70], 8.3)

heading(doc, "Fortalezas verificadas", 2)
bullets(doc, [
    "Validacion real del token con Supabase Auth en el servidor; los perfiles inexistentes o inactivos fallan cerrados.",
    "MFA AAL2 exigido para administradores y perfiles con captura antes de ejecutar mutaciones.",
    "CAPTCHA Turnstile integrado al inicio de sesion y registro publico deshabilitado segun configuracion reportada por administracion.",
    "Filtro por sucursal, marca y gerente implementado en servidor para funnel_records.",
    "RLS habilitado y privilegios directos revocados a anon y authenticated; la escritura se centraliza en la API.",
    "Contraseñas nuevas con minimo de 12 caracteres y complejidad; secretos de servicio solo referenciados en servidor.",
    "CSP, HSTS, anti-framing, nosniff, politica de permisos, no-store en API y defensa de origen cruzado.",
    "Bitacora de eventos sensibles y alertas webhook para borrados, reemplazos e importaciones parciales.",
    "Dependabot semanal y flujo de CI con auditoria de dependencias y build de produccion.",
])

heading(doc, "Decisiones inmediatas", 2)
add_table(doc, ["Prioridad", "Accion", "Plazo recomendado"], [
    ["P0", "Actualizar Next.js a >=16.2.11 y sustituir/actualizar xlsx a una version corregida; actualizar sharp/postcss mediante el arbol compatible.", "24-72 horas"],
    ["P0", "Cerrar la API de almacenamiento con lista permitida de claves y autorizacion de lectura/escritura por tipo de dato.", "7 dias"],
    ["P0", "Aplicar en servidor los permisos de captura por campo (captureFields), no solo en la interfaz.", "7 dias"],
    ["P1", "Cambiar el trigger de perfil nuevo a active=false y activacion solo por administrador.", "7 dias"],
    ["P1", "Probar alertas, limitacion de tasa, respaldo y restauracion documentada.", "30 dias"],
], [0.65, 4.65, 1.10], 8.3)

heading(doc, "2. Alcance, metodologia y limitaciones", 1)
para(doc, "Alcance revisado: aplicacion Next.js, endpoint /api/storage, integracion con Supabase Auth y base de datos, esquema SQL, configuracion de cabeceras, dependencias npm, automatizacion de GitHub, manejo de usuarios, importaciones y mecanismos de auditoria.")
para(doc, "Metodologia: revision estatica inspirada en OWASP ASVS, principio de minimo privilegio, defensa en profundidad, confidencialidad-integridad-disponibilidad y revision de configuracion. Se ejecuto una auditoria de dependencias contra el registro npm el 31/julio/2026.")
add_callout(doc, "Limite de la conclusion", "No se realizaron pruebas de penetracion, escaneo dinamico del despliegue, revision directa de paneles productivos, prueba de restauracion, examen forense de logs ni validacion independiente de Vercel/Supabase. Por ello, el reporte expresa seguridad observada y riesgos, no una garantia contra ataques.", "medium")

heading(doc, "Activos y datos protegidos", 2)
add_table(doc, ["Activo", "Clasificacion", "Riesgo principal"], [
    ["Credenciales, sesiones y secretos", "Critico", "Toma de cuenta o control total de la base si se expone service role."],
    ["Perfiles, roles y alcances", "Alto", "Escalamiento de privilegios o acceso a sucursales no autorizadas."],
    ["Funnel, BDC y campañas", "Alto", "Exposicion o alteracion de resultados comerciales."],
    ["Industria, AMDA y geointeligencia", "Alto", "Divulgacion de estrategia, ventas y ubicaciones."],
    ["Bitacora de auditoria", "Alto", "Perdida de trazabilidad o encubrimiento de actividad."],
    ["Exportaciones PDF/Excel/PPT", "Alto", "Copia fuera de la plataforma sin controles posteriores."],
], [2.05, 1.15, 3.20], 8.4)

heading(doc, "3. Arquitectura y fronteras de confianza", 1)
add_table(doc, ["Capa", "Funcion", "Confianza / control"], [
    ["Navegador", "Interfaz, filtros, importacion y exportacion", "No confiable; cualquier usuario puede modificar solicitudes manualmente."],
    ["Cloudflare Turnstile", "Validacion anti-bot en autenticacion", "Control externo; debe comprobarse dominio y secreto."],
    ["Supabase Auth", "Credenciales, sesion, MFA y claims", "Fuente de identidad; el token se revalida en servidor."],
    ["Next.js / Vercel", "API, autorizacion, validacion y auditoria", "Frontera central; usa service role y debe aplicar todo alcance."],
    ["Supabase Postgres", "Perfiles, registros, almacenamiento y auditoria", "RLS activo; service role omite RLS por diseño."],
    ["Webhook de alertas", "Notificacion operativa", "Canal externo; estado de entrega no verificado."],
], [1.35, 2.25, 2.80], 8.3)
para(doc, "Flujo principal: navegador -> Supabase Auth/Turnstile -> token Bearer -> API Next.js -> validacion de perfil y AAL -> operacion con service role -> Supabase -> auditoria y, cuando aplica, webhook.")

heading(doc, "4. Evaluacion detallada de controles", 1)
controls = [
    ["Autenticacion", "Implementado", "getUser(token), perfil activo, CAPTCHA y MFA AAL2 para captura/administracion."],
    ["Autorizacion de funnel_records", "Implementado", "Alcance por sucursal, marca derivada y gerente en consultas y mutaciones."],
    ["Autorizacion de almacenamiento compartido", "Parcial", "GET generico devuelve documentos completos; falta mapa de lectura por clave y alcance."],
    ["Permisos por campo", "No implementado en API", "captureFields se utiliza en interfaz, pero no se valida cada metrica del payload en servidor."],
    ["Alta de cuentas", "Parcial", "Creacion administrativa fuerte; trigger SQL crea perfil capturista activo ante usuario Auth nuevo."],
    ["Base de datos", "Implementado", "RLS activo y privilegios directos revocados; funciones y service role requieren gobierno estricto."],
    ["Validacion de entradas", "Parcial", "Limites de tamaño, periodo y longitud; faltan rangos numericos y validacion semantica exhaustiva."],
    ["Proteccion del navegador", "Implementado con excepcion", "Cabeceras completas; CSP conserva unsafe-inline para scripts y estilos."],
    ["Auditoria", "Parcial", "Eventos sensibles registrados; sin politica de retencion/inmutabilidad demostrada."],
    ["Alertamiento", "Parcial", "Webhook programado y resumen admin; prueba de entrega y escalamiento no evidenciada."],
    ["Dependencias", "Deficiente al corte", "15 vulnerabilidades: 9 altas y 6 moderadas."],
    ["Respaldo/recuperacion", "No demostrado", "Exportacion manual disponible; no hay restauracion de prueba ni RTO/RPO."],
]
add_table(doc, ["Control", "Estado", "Evidencia / observacion"], controls, [1.60, 1.35, 3.45], 8.1)

heading(doc, "Cabeceras verificadas", 2)
bullets(doc, [
    "Content-Security-Policy con default-src self, object-src none, frame-ancestors none y form-action self.",
    "Strict-Transport-Security por dos años, includeSubDomains y preload.",
    "X-Content-Type-Options nosniff, X-Frame-Options DENY y Referrer-Policy strict-origin-when-cross-origin.",
    "Permissions-Policy desactiva camara, microfono, geolocalizacion, pagos y USB.",
    "API marcada Cache-Control: no-store; COOP y CORP en same-origin.",
])

heading(doc, "5. Hallazgos y riesgos residuales", 1)
risks = [
    ["R-01", "Critico", "Almacenamiento compartido sin lista cerrada", "Un usuario autenticado puede consultar claves genericas completas y, para claves no clasificadas, escribir almacenamiento shared. Geo/cutoff/backup no aparecen en conjuntos protegidos.", "Definir allowlist por clave, rol, accion y alcance; rechazar por defecto."],
    ["R-02", "Alto", "Permisos de captura por campo solo en interfaz", "Un perfil personalizado podria enviar por API metricas no autorizadas aunque la UI las oculte.", "Validar captureFields y limpiar/rechazar payload en servidor."],
    ["R-03", "Alto", "Dependencias vulnerables", "9 altas y 6 moderadas: Next.js <16.2.11, xlsx 0.18.5, sharp y postcss transitivos.", "Actualizar, reconstruir lockfile, probar y exigir audit sin altas en CI."],
    ["R-04", "Alto", "Recuperacion no probada", "No hay evidencia de respaldo restaurable, RTO/RPO ni simulacro.", "Definir RPO/RTO; respaldos cifrados; restauracion trimestral."],
    ["R-05", "Alto", "Alta automatica activa", "El trigger de auth.users crea perfil capturista active=true si se habilita signup por error.", "Crear perfil inactivo y activarlo solo por administrador; alerta por alta."],
    ["R-06", "Medio-alto", "Abuso y limitacion de tasa no demostrados", "CAPTCHA protege login, pero no se comprobo rate limit en API de importacion/almacenamiento.", "Rate limits por usuario/IP, cuotas por accion y respuestas 429."],
    ["R-07", "Medio-alto", "Service role como punto critico", "La API opera con privilegio amplio y omite RLS; una fuga o falla de autorizacion afecta toda la base.", "Rotacion, minimo uso, monitoreo y RPCs con permisos especificos."],
    ["R-08", "Medio", "Auditoria alterable y sin retencion", "No se demostro inmutabilidad, archivado, limpieza controlada ni retencion.", "Politica 12-24 meses, append-only y exportacion a repositorio separado."],
    ["R-09", "Medio", "Alertas sin prueba operativa", "Existe codigo de webhook, pero no evidencia de prueba, receptor, guardias o SLA.", "Prueba mensual, runbook, dueño y confirmacion de entrega."],
    ["R-10", "Medio", "Validacion de importaciones incompleta", "Hay limites basicos, no rangos de KPI, conteo total robusto, ni pruebas de archivos hostiles.", "Esquema estricto, rangos, limites de celdas/tiempo y pruebas maliciosas."],
    ["R-11", "Medio", "CSP permite unsafe-inline", "Reduce la proteccion contra XSS si aparece una inyeccion.", "Migrar a nonces/hashes y eliminar unsafe-inline gradualmente."],
    ["R-12", "Medio", "Sin pruebas negativas automatizadas", "CI compila y audita, pero no demuestra que un usuario de sucursal A no lea/escriba B.", "Suite de autorizacion por rol, sucursal, marca, gerente y campo."],
    ["R-13", "Medio", "Gobierno de exportaciones", "Los archivos exportados quedan fuera del control de acceso de la plataforma.", "Marca de agua, clasificacion, minimo dato y politica de custodia."],
    ["R-14", "Medio", "Respuesta a incidentes no formalizada", "No hay evidencia de matriz de contactos, severidades, contencion y comunicacion.", "Runbook, simulacro semestral y registro de incidentes."],
]
add_table(doc, ["ID", "Nivel", "Hallazgo", "Escenario / impacto", "Tratamiento"], risks, [0.42, 0.63, 1.35, 2.20, 1.80], 7.2)

heading(doc, "6. Auditoria de dependencias al corte", 1)
add_callout(doc, "Resultado pnpm audit", "15 vulnerabilidades encontradas: 9 altas y 6 moderadas. La ejecucion se realizo el 31/julio/2026 contra el registro oficial de npm. El resultado puede cambiar con nuevas publicaciones.", "high")
deps = [
    ["next", "16.2.6", "4 altas", "Bypass middleware/proxy, DoS y SSRF", ">=16.2.11"],
    ["xlsx", "0.18.5", "2 altas", "Prototype pollution y ReDoS", ">=0.20.2 segun advisories; validar canal oficial/licencia"],
    ["sharp (transitiva)", "via next", "1 alta agrupada", "Vulnerabilidades heredadas de libvips", ">=0.35.0"],
    ["postcss (transitiva)", "via next", "2 altas", "Lectura de archivos/path traversal por source maps", ">=8.5.18"],
    ["otras", "varias", "6 moderadas", "Revisar salida completa tras actualizar", "Resolver o aceptar formalmente"],
]
add_table(doc, ["Paquete", "Version/ruta", "Severidad", "Riesgo", "Remediacion"], deps, [1.25, 0.90, 0.85, 1.85, 1.55], 7.8)
para(doc, "Criterio de salida: no desplegar una nueva version con vulnerabilidades altas conocidas sin excepcion documentada, analisis de explotabilidad, responsable y fecha de caducidad. Ejecutar nuevamente pnpm audit y pruebas funcionales despues de actualizar.")

heading(doc, "7. Plan de tratamiento 0-90 dias", 1)
add_table(doc, ["Horizonte", "Acciones", "Evidencia de cierre"], [
    ["0-7 dias", "Actualizar dependencias; allowlist de storage; permisos por campo en API; trigger active=false; comprobar que signup sigue deshabilitado.", "Audit sin altas; pruebas negativas; SQL aplicado; capturas/config exportada."],
    ["8-30 dias", "Rate limiting; prueba webhook; rotacion de secretos; inventario de datos; politica de exportaciones; runbook de incidentes.", "Pruebas 429; alerta recibida; acta de rotacion; runbook aprobado."],
    ["31-60 dias", "Respaldo y restauracion; retencion de auditoria; recertificacion de accesos; pruebas de importacion hostil.", "Restauracion firmada; consultas de retencion; matriz de usuarios; resultados de QA."],
    ["61-90 dias", "Pentest externo; remediar CSP; threat model; monitoreo central; simulacro de incidente.", "Informe pentest y retest; CSP sin inline; diagrama; informe de simulacro."],
], [1.0, 3.65, 1.75], 8.2)

heading(doc, "Responsabilidades sugeridas", 2)
add_table(doc, ["Rol", "Responsabilidad"], [
    ["Propietario de negocio", "Aprobar riesgos, datos, RTO/RPO y perfiles."],
    ["Administrador de plataforma", "Altas/bajas, MFA, recertificacion, alertas y exportaciones."],
    ["Desarrollo", "Autorizacion, actualizaciones, pruebas y correcciones."],
    ["Infraestructura/Seguridad", "Secretos, WAF, respaldos, monitoreo y respuesta."],
    ["Auditoria", "Revisar evidencias, excepciones y cierre de hallazgos."],
], [1.85, 4.55], 8.5)

heading(doc, "8. Controles gratuitos, de pago y migracion interna", 1)
add_table(doc, ["Capacidad", "Puede hacerse ahora", "Licencia / observacion"], [
    ["MFA TOTP, CAPTCHA, RLS, validacion y cabeceras", "Si", "Implementable con codigo/configuracion; verificar limites vigentes del proveedor."],
    ["Auditoria propia y webhook", "Si", "Coste depende del receptor y retencion."],
    ["Dependabot/CI y pruebas", "Si", "GitHub puede imponer limites segun plan/repositorio."],
    ["Respaldo manual cifrado", "Si", "Requiere disciplina y prueba de restauracion."],
    ["PITR, respaldo administrado y mayor retencion", "Segun plan", "Recomendable Supabase de pago para produccion; confirmar oferta vigente."],
    ["WAF/rate limiting administrado, equipo y uso comercial", "Segun plan", "Vercel Pro o servicio equivalente; Hobby no debe asumirse como nivel empresarial."],
    ["Pentest, SIEM, guardias y continuidad", "No incluidos", "Servicios/procesos adicionales, independientemente del hosting."],
], [2.45, 1.30, 2.65], 8.1)

heading(doc, "Migracion a servidor interno", 2)
para(doc, "Migrar no elimina los riesgos: transfiere a la empresa la responsabilidad de seguridad fisica/logica, parcheo, alta disponibilidad, monitoreo, copias y recuperacion. Antes de mover produccion se recomienda:")
bullets(doc, [
    "Separar Internet, proxy/WAF, aplicacion y base de datos en redes distintas; Postgres sin exposicion publica.",
    "TLS administrado y renovacion automatica; secretos en vault, nunca en repositorio ni archivos compartidos.",
    "Cuenta de servicio minima, rotacion y registros centralizados; EDR y parcheo mensual del sistema operativo.",
    "Respaldos cifrados 3-2-1, copia fuera del sitio, RPO/RTO aprobados y restauracion trimestral.",
    "Monitoreo de disponibilidad, errores 5xx, accesos denegados, cambios de privilegio, exportaciones y borrados.",
    "Staging separado, despliegue reproducible, rollback y prueba de carga/capacidad.",
    "SSO/MFA corporativo cuando sea posible, revisiones trimestrales de acceso y baja inmediata de personal.",
])

heading(doc, "9. Plan de pruebas de seguridad", 1)
tests = [
    ["AUTH-01", "Usuario inactivo", "Login rechazado y evento registrado."],
    ["AUTH-02", "Admin/capturista sin AAL2 intenta POST/DELETE", "HTTP 403; ninguna modificacion."],
    ["AUTHZ-01", "Usuario de Juventud consulta/escribe otra sucursal", "Solo datos permitidos o 403; prueba en records y storage."],
    ["AUTHZ-02", "Perfil limitado a SOLICITUDES envia VENTAS", "Servidor rechaza o elimina campo no permitido."],
    ["AUTHZ-03", "Viewer usa clave shared desconocida", "HTTP 400/403 por clave no permitida."],
    ["IMPORT-01", "Archivo > limite, periodo invalido, valores negativos/extremos", "Rechazo previo a escritura y auditoria."],
    ["ABUSE-01", "Rafaga de login/API/importacion", "CAPTCHA/429/bloqueo temporal; alerta si corresponde."],
    ["AUDIT-01", "Borrado permanente e importacion parcial", "Evento completo y webhook recibido."],
    ["DR-01", "Restauracion a ambiente aislado", "RPO/RTO medidos, integridad validada y acta firmada."],
    ["SESSION-01", "Logout, token expirado y revocado", "Sesion invalidada sin pantalla de error ni reutilizacion."],
]
add_table(doc, ["ID", "Prueba", "Resultado esperado"], tests, [0.75, 3.15, 2.50], 8.2)

heading(doc, "10. Evidencia tecnica revisada", 1)
add_table(doc, ["Archivo / fuente", "Evidencia principal"], [
    ["app/api/storage/route.ts", "Autenticacion, MFA, alcance records, validaciones, auditoria, webhook, CRUD usuarios y riesgos de claves genericas/campos."],
    ["lib/supabase.ts", "Validacion token, lectura de claims y perfil activo; service role solo servidor."],
    ["supabase/schema.sql", "Tablas, RLS, revocacion, triggers, funciones, perfiles y auditoria."],
    ["next.config.ts", "Cabeceras y no-store de API."],
    ["app/page.tsx", "Turnstile, sesion, MFA, importacion, permisos de interfaz y exportaciones."],
    ["package.json / pnpm audit", "Versiones y 15 vulnerabilidades al corte."],
    [".github/dependabot.yml y workflow", "Actualizacion semanal, audit y build."],
    [".gitignore y .env.example", "Archivos locales de secretos excluidos; no se detectaron valores de service role en archivos revisados."],
], [2.35, 4.05], 8.2)

heading(doc, "11. Referencias", 1)
refs = [
    "OWASP Application Security Verification Standard (ASVS): https://owasp.org/www-project-application-security-verification-standard/",
    "OWASP Authentication Cheat Sheet: https://cheatsheetseries.owasp.org/cheatsheets/Authentication_Cheat_Sheet.html",
    "OWASP Logging Cheat Sheet: https://cheatsheetseries.owasp.org/cheatsheets/Logging_Cheat_Sheet.html",
    "Supabase production checklist: https://supabase.com/docs/guides/deployment/going-into-prod",
    "Supabase password security: https://supabase.com/docs/guides/auth/password-security",
    "Supabase CAPTCHA: https://supabase.com/docs/guides/auth/auth-captcha",
    "Vercel Hobby plan: https://vercel.com/docs/plans/hobby",
    "Vercel rate limiting: https://vercel.com/docs/vercel-firewall/vercel-waf/rate-limiting",
    "GitHub advisories observados: GHSA-4r6h-8v6p-xvw6, GHSA-5pgg-2g8v-p4x9, GHSA-6gpp-xcg3-4w24, GHSA-m99w-x7hq-7vfj, GHSA-89xv-2m56-2m9x, GHSA-p9j2-gv94-2wf4, GHSA-6g55-p6wh-862q y GHSA-r28c-9q8g-f849.",
]
bullets(doc, refs)

heading(doc, "12. Conclusion", 1)
para(doc, "La plataforma ya incorpora controles valiosos y no parte de cero. Su mayor fortaleza es que la autenticacion y parte importante de la autorizacion viven en el servidor, respaldadas por Supabase Auth, MFA, CAPTCHA y una base con acceso directo revocado. Su mayor debilidad actual es la inconsistencia entre el control estricto de funnel_records y el almacenamiento JSON compartido, junto con permisos por campo que aun dependen de la interfaz y dependencias vulnerables al corte.")
para(doc, "La recomendacion es mantener el servicio operativo solo con acceso controlado mientras se ejecutan las acciones P0, evitar ampliar usuarios o datos sensibles hasta cerrar R-01 a R-05 y no comunicar que la plataforma es 'imposible de hackear'. La declaracion defendible es: cuenta con controles de seguridad en capas, se audita periodicamente y mantiene un plan formal de tratamiento de riesgos.")
add_callout(doc, "Criterio para elevar la postura a riesgo MEDIO", "Cerrar R-01, R-02, R-03 y R-05; demostrar rate limiting, webhook operativo y pruebas negativas; aprobar y ejecutar una restauracion. Despues, realizar un pentest y retest independiente.", "ok")

# Metadata
doc.core_properties.title = "Reporte integral de seguridad - Funnel Comercial CR3"
doc.core_properties.subject = "Revision de seguridad de aplicacion, datos y operacion"
doc.core_properties.author = "Grupo CR3"
doc.core_properties.keywords = "seguridad, Supabase, Vercel, Funnel, auditoria, riesgo"
doc.core_properties.comments = "Documento confidencial de uso interno"

doc.save(DOCX)
print(DOCX)
